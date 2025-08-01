from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os
from django.conf import settings
from students.models import Student, StudentTestResult
from schools.models import School, Class
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
import base64

def create_table_border(table):
    """Add borders to table"""
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tbl.tblPr.append(tblBorders)

def generate_individual_report(student_id):
    """Generate individual student report"""
    student = Student.objects.get(id=student_id)
    doc = Document()

    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"{student.class_assigned.school.name} - Physical Assessment Report"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    title = doc.add_heading('PHYSICAL ASSESSMENT REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Student Information
    doc.add_heading('Student Information', level=1)

    # Create student info table
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = 'Table Grid'
    create_table_border(info_table)

    # Fill student information
    info_data = [
        ('Name', student.name),
        ('Roll Number', student.roll_number),
        ('Class', f"Grade {student.class_assigned.grade} Section {student.class_assigned.section}"),
        ('Date of Birth', student.date_of_birth.strftime('%Y-%m-%d')),
        ('Age', f"{student.age} years"),
        ('Gender', student.get_gender_display()),
        ('Height', f"{student.height} inches"),
        ('Weight', f"{student.weight} kg"),
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    # BMI Analysis
    doc.add_heading('BMI Analysis', level=1)
    bmi_para = doc.add_paragraph()
    bmi_para.add_run(f"BMI: {student.bmi}").bold = True
    bmi_para.add_run(f" - Category: {student.get_bmi_category_display()}")

    # Physical Test Results
    test_results = student.test_results.all()
    if test_results:
        doc.add_heading('Physical Test Results', level=1)

        test_table = doc.add_table(rows=len(test_results) + 1, cols=4)
        test_table.style = 'Table Grid'
        create_table_border(test_table)

        # Headers
        headers = ['Test Name', 'Score', 'Unit', 'Comment']
        for i, header in enumerate(headers):
            test_table.rows[0].cells[i].text = header
            test_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

        # Fill test data
        for i, result in enumerate(test_results, 1):
            test_table.rows[i].cells[0].text = result.test.name
            test_table.rows[i].cells[1].text = str(result.score)
            test_table.rows[i].cells[2].text = result.test.unit
            test_table.rows[i].cells[3].text = result.comment or 'N/A'

    # Overall Comment
    if student.overall_comment:
        doc.add_heading('Overall Assessment', level=1)
        doc.add_paragraph(student.overall_comment)

    # Save document
    filename = f"report_{student.name.replace(' ', '_')}_{student.roll_number}.docx"
    file_path = os.path.join(settings.MEDIA_ROOT, 'reports', filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)

    return file_path, filename

def generate_class_report(class_id):
    """Generate class-wise report"""
    from schools.models import Class
    class_obj = Class.objects.get(id=class_id)
    students = class_obj.students.all()

    doc = Document()

    # Title
    title = doc.add_heading(f'CLASS REPORT - Grade {class_obj.grade} Section {class_obj.section}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # School info
    doc.add_paragraph(f"School: {class_obj.school.name}")
    doc.add_paragraph(f"Total Students: {students.count()}")
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # Class Summary Table
    doc.add_heading('Class Summary', level=1)

    summary_table = doc.add_table(rows=len(students) + 1, cols=8)
    summary_table.style = 'Table Grid'
    create_table_border(summary_table)

    # Headers
    headers = ['Roll', 'Name', 'Age', 'Gender', 'Height', 'Weight', 'BMI', 'BMI Category']
    for i, header in enumerate(headers):
        summary_table.rows[0].cells[i].text = header
        summary_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    # Fill student data
    for i, student in enumerate(students, 1):
        row_data = [
            student.roll_number,
            student.name,
            str(student.age),
            student.get_gender_display(),
            f"{student.height}\"",
            f"{student.weight}kg",
            str(student.bmi),
            student.get_bmi_category_display()
        ]

        for j, data in enumerate(row_data):
            summary_table.rows[i].cells[j].text = data

    # BMI Distribution
    doc.add_heading('BMI Distribution', level=1)
    bmi_categories = students.values('bmi_category').distinct()
    for category in bmi_categories:
        count = students.filter(bmi_category=category['bmi_category']).count()
        doc.add_paragraph(f"{category['bmi_category'].title()}: {count} students")

    # Save document
    filename = f"class_report_grade_{class_obj.grade}_section_{class_obj.section}.docx"
    file_path = os.path.join(settings.MEDIA_ROOT, 'reports', filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)

    return file_path, filename

def generate_school_report(school_id):
    """Generate school-wide report"""
    school = School.objects.get(id=school_id)
    students = Student.objects.filter(class_assigned__school=school)

    doc = Document()

    # Title
    title = doc.add_heading(f'SCHOOL PHYSICAL ASSESSMENT REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # School info
    doc.add_paragraph(f"School: {school.name}")
    doc.add_paragraph(f"Address: {school.address}")
    doc.add_paragraph(f"Total Students: {students.count()}")
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # Overall Statistics
    doc.add_heading('Overall Statistics', level=1)

    # BMI Distribution
    bmi_stats = {}
    for category in ['underweight', 'healthy', 'overweight', 'obese']:
        count = students.filter(bmi_category=category).count()
        percentage = (count / students.count() * 100) if students.count() > 0 else 0
        bmi_stats[category] = {'count': count, 'percentage': percentage}

    stats_table = doc.add_table(rows=5, cols=3)
    stats_table.style = 'Table Grid'
    create_table_border(stats_table)

    # Headers
    stats_table.rows[0].cells[0].text = "BMI Category"
    stats_table.rows[0].cells[1].text = "Count"
    stats_table.rows[0].cells[2].text = "Percentage"

    for i, (category, data) in enumerate(bmi_stats.items(), 1):
        stats_table.rows[i].cells[0].text = category.title()
        stats_table.rows[i].cells[1].text = str(data['count'])
        stats_table.rows[i].cells[2].text = f"{data['percentage']:.1f}%"

    # Class-wise breakdown
    doc.add_heading('Class-wise Breakdown', level=1)

    classes = school.classes.all()
    class_table = doc.add_table(rows=len(classes) + 1, cols=5)
    class_table.style = 'Table Grid'
    create_table_border(class_table)

    # Headers
    class_headers = ['Class', 'Total Students', 'Avg BMI', 'Healthy %', 'At Risk %']
    for i, header in enumerate(class_headers):
        class_table.rows[0].cells[i].text = header
        class_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    for i, class_obj in enumerate(classes, 1):
        class_students = class_obj.students.all()
        total = class_students.count()
        avg_bmi = sum([s.bmi for s in class_students]) / total if total > 0 else 0
        healthy_count = class_students.filter(bmi_category='healthy').count()
        at_risk_count = class_students.exclude(bmi_category='healthy').count()

        class_table.rows[i].cells[0].text = f"Grade {class_obj.grade} Sec {class_obj.section}"
        class_table.rows[i].cells[1].text = str(total)
        class_table.rows[i].cells[2].text = f"{avg_bmi:.1f}"
        class_table.rows[i].cells[3].text = f"{(healthy_count/total*100) if total > 0 else 0:.1f}%"
        class_table.rows[i].cells[4].text = f"{(at_risk_count/total*100) if total > 0 else 0:.1f}%"

    # Save document
    filename = f"school_report_{school.name.replace(' ', '_')}.docx"
    file_path = os.path.join(settings.MEDIA_ROOT, 'reports', filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)

    return file_path, filename

# Apply K-means clustering
def apply_clustering():
    """Apply K-means clustering to group students by performance"""
    from sklearn.cluster import KMeans
    import numpy as np

    students = Student.objects.all()
    if students.count() < 3:
        return False

    # Prepare features for clustering
    features = []
    student_ids = []

    for student in students:
        # Use BMI, height, weight as features
        features.append([student.bmi, student.height, student.weight])
        student_ids.append(student.id)

    features = np.array(features)

    # Apply K-means clustering
    n_clusters = min(3, len(features))  # 3 groups max
    kmeans = KMeans(n_clusters=n_clusters, random_state=42)
    clusters = kmeans.fit_predict(features)

    # Update students with cluster groups
    for i, student_id in enumerate(student_ids):
        student = Student.objects.get(id=student_id)
        student.performance_group = clusters[i] + 1  # 1-based indexing
        student.save()

    return True
